# comment_inserter.py

from docx import Document

def insert_comment(doc_path, target_clause, comment_text, output_path):
    """
    Inserts a comment or marked text next to the target clause in a .docx file.

    Args:
        doc_path (str): Path to the input .docx file
        target_clause (str): The clause or sentence to annotate
        comment_text (str): The feedback or compliance suggestion
        output_path (str): Path to save the annotated output .docx
    """
    doc = Document(doc_path)
    found = False

    for para in doc.paragraphs:
        if target_clause.strip() in para.text:
            run = para.add_run(f"\n⚠️ COMMENT: {comment_text}")
            run.bold = True
            run.italic = True
            found = True
            break

    if not found:
        print(f"⚠️ Warning: Clause not found: '{target_clause[:40]}...'")

    doc.save(output_path)
    print(f"✅ Comment inserted and saved to {output_path}")
